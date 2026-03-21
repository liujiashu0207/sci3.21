#!/usr/bin/env python3
"""RDA* 论文修改稿 v3 — 整合全部修订"""
import os
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()
for s in doc.sections:
    s.page_width=Cm(21); s.page_height=Cm(29.7)
    s.top_margin=Cm(2.5); s.bottom_margin=Cm(2.5)
    s.left_margin=Cm(2.8); s.right_margin=Cm(2.8)
st = doc.styles['Normal']
st.font.name='宋体'; st.font.size=Pt(10.5)
st.element.rPr.rFonts.set(qn('w:eastAsia'),'宋体')
st.paragraph_format.line_spacing=1.25; st.paragraph_format.space_after=Pt(0)

def heading(text, level=1):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER if level==0 else WD_ALIGN_PARAGRAPH.LEFT
    r=p.add_run(text); r.bold=True
    r.font.size={0:Pt(16),1:Pt(14),2:Pt(12),3:Pt(11)}[level]
    r.font.name='黑体'; r.element.rPr.rFonts.set(qn('w:eastAsia'),'黑体')
    p.paragraph_format.space_before=Pt(12); p.paragraph_format.space_after=Pt(6)

def para(text, sz=10.5, font='宋体', indent=True, bold=False):
    p=doc.add_paragraph()
    if indent: p.paragraph_format.first_line_indent=Cm(0.74)
    r=p.add_run(text); r.font.size=Pt(sz); r.font.name=font
    r.element.rPr.rFonts.set(qn('w:eastAsia'),font); r.bold=bold; return p

def formula(text, num=''):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(f'{text}　　{num}'); r.font.name='Times New Roman'; r.font.size=Pt(10.5); r.italic=True

def caption(cn, en):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(cn); r.bold=True; r.font.size=Pt(9); r.font.name='黑体'; r.element.rPr.rFonts.set(qn('w:eastAsia'),'黑体')
    p2=doc.add_paragraph(); p2.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r2=p2.add_run(en); r2.bold=True; r2.font.size=Pt(9); r2.font.name='Times New Roman'

def tbl(headers, rows, footnote=None):
    t=doc.add_table(rows=1+len(rows), cols=len(headers)); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.style='Table Grid'
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; c.text=''; p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        r=p.add_run(h); r.bold=True; r.font.size=Pt(9); r.font.name='Times New Roman'; r.element.rPr.rFonts.set(qn('w:eastAsia'),'宋体')
    for ri,row in enumerate(rows):
        for ci,val in enumerate(row):
            c=t.rows[ri+1].cells[ci]; c.text=''; p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
            r=p.add_run(str(val)); r.font.size=Pt(9); r.font.name='Times New Roman'; r.element.rPr.rFonts.set(qn('w:eastAsia'),'宋体')
            if 'RDA*' in str(val) or '本文' in str(val): r.bold=True
    if footnote:
        p=doc.add_paragraph(); p.paragraph_format.first_line_indent=Cm(0)
        r=p.add_run(footnote); r.font.size=Pt(8); r.font.name='宋体'; r.element.rPr.rFonts.set(qn('w:eastAsia'),'宋体')

def img(path, w=14):
    if os.path.exists(path):
        doc.add_picture(path, width=Cm(w)); doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER

def bpara(label, text, lf='黑体'):
    p=doc.add_paragraph(); p.paragraph_format.first_line_indent=Cm(0.74)
    r1=p.add_run(label); r1.bold=True; r1.font.size=Pt(10.5); r1.font.name=lf; r1.element.rPr.rFonts.set(qn('w:eastAsia'),lf)
    r2=p.add_run(text); r2.font.size=Pt(10.5); r2.font.name='宋体'; r2.element.rPr.rFonts.set(qn('w:eastAsia'),'宋体')

# ═════════════ 标题+作者 ═════════════
heading('残差驱动自适应加权A*路径规划算法', 0)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('刘家树'); r.font.size=Pt(12); r.font.name='楷体'; r.element.rPr.rFonts.set(qn('w:eastAsia'),'楷体')
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('南方科技大学 机械与能源工程系，广东 深圳 518055'); r.font.size=Pt(9)

# ═════════════ 中文摘要 (r_res符号) ═════════════
bpara('摘　要：',
    '针对A*算法在复杂栅格环境中搜索效率受限、路径质量不佳的问题，提出了一种残差驱动自适应加权A*路径规划算法'
    '（Residual-Driven Adaptive Weighted A*，RDA*）。该算法以Octile距离作为启发函数，引入局部障碍物密度概念，'
    '通过积分图实现O(1)复杂度的逐节点障碍率查询，并定义通行残差r_res(n)=1−ρ_local(n)作为权重调节的驱动因子，'
    '设计自适应权重函数α(n)=1+β·r_res(n)，使搜索在空旷区域加速探索、在密集区域保守寻优；同时采用两阶段路径平滑'
    '策略（视线简化与角点插值），在保证路径合法性的前提下显著缩短路径长度并减少转弯次数。在MovingAI公开基准数据集'
    '的15张地图上，以20条任务×7种算法进行对比实验，结果表明：与A*(欧氏)相比，RDA*的路径长度缩短3.6%（p<0.001），'
    '转弯次数减少76.4%（p<0.001），扩展节点减少61.5%（p<0.001），总运行时间缩短31.9%；消融实验验证了自适应权重'
    '与路径平滑两个模块各自的独立贡献。补充实验在长路径场景下进一步验证了算法的鲁棒性。')
bpara('关键词：', '路径规划；A*算法；自适应权重；障碍物密度；路径平滑；栅格地图')
bpara('文献标志码：', 'A　中图分类号：TP242')

# 英文摘要 (r_res符号)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(12)
r=p.add_run('Residual-Driven Adaptive Weighted A* Path Planning Algorithm'); r.bold=True; r.font.size=Pt(12); r.font.name='Times New Roman'
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('LIU Jiashu'); r.font.size=Pt(10); r.font.name='Times New Roman'
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('Department of Mechanical and Energy Engineering, Southern University of Science and Technology, Shenzhen, Guangdong 518055, China'); r.font.size=Pt(9); r.font.name='Times New Roman'

p=doc.add_paragraph(); p.paragraph_format.first_line_indent=Cm(0.74)
r=p.add_run('Abstract: '); r.bold=True; r.font.size=Pt(9); r.font.name='Times New Roman'
r2=p.add_run(
    'To address the limited search efficiency and suboptimal path quality of the A* algorithm in complex grid environments, '
    'a Residual-Driven Adaptive Weighted A* (RDA*) path planning algorithm is proposed. The algorithm employs octile distance '
    'as the heuristic function and introduces a local obstacle density concept, achieving O(1) per-node obstacle ratio queries '
    'via integral images. A traversal residual r_res(n)=1-rho_local(n) is defined as the driving factor for weight adjustment, '
    'and an adaptive weight function alpha(n)=1+beta*r_res(n) is designed to accelerate exploration in open areas while '
    'maintaining conservative search in dense regions. A two-stage path smoothing strategy significantly reduces path length '
    'and turn count while ensuring path legality. Experiments on 15 maps from the MovingAI benchmark dataset show that compared '
    'with A*(Euclidean), RDA* reduces path length by 3.6% (p<0.001), turn count by 76.4% (p<0.001), expanded nodes by 61.5% '
    '(p<0.001), and total runtime by 31.9%. Ablation studies confirm the independent contributions of the adaptive weight and '
    'path smoothing modules.')
r2.font.size=Pt(9); r2.font.name='Times New Roman'
p=doc.add_paragraph(); p.paragraph_format.first_line_indent=Cm(0.74)
r=p.add_run('Keywords: '); r.bold=True; r.font.size=Pt(9); r.font.name='Times New Roman'
r2=p.add_run('path planning; A* algorithm; adaptive weight; obstacle density; path smoothing; grid map'); r2.font.size=Pt(9); r2.font.name='Times New Roman'

# ═════════════ 0 引言 (引用序号按首次引用顺序) ═════════════
heading('0　引言', 1)
para('路径规划是实现移动机器人自主导航的关键技术之一，其目标是在已知或部分已知的环境中规划一条从起点到目标点的'
     '最优路径，同时避开障碍物并满足路径长度、转弯次数、搜索效率等多项性能指标[1]。随着机器人技术在物流仓储[2]、'
     '巡检作业[3]、游戏AI[4]等领域的广泛应用，路径规划算法的性能需求日益提高。')
para('A*算法作为一种经典的启发式搜索算法[5]，通过代价函数f(n)=g(n)+h(n)引导搜索方向，兼顾了Dijkstra算法的最优性'
     '和贪心搜索的高效性，是目前全局路径规划中应用最为广泛的算法之一。然而，传统A*算法在复杂栅格环境中仍存在'
     '以下不足：（1）搜索效率受限，在大规模地图中扩展节点过多，导致运行时间较长；（2）规划路径包含大量冗余转折点，'
     '不利于机器人平稳运行；（3）启发函数权重固定，无法根据局部环境特征自适应调整搜索策略。')
para('针对上述问题，国内外学者进行了大量研究。在启发函数权重优化方面，Algfoor等[6]提出三种加权启发式搜索技术，'
     '分别基于迭代次数、起终点距离和行进代价设计权重，并应用于A*、双向A*和JPS算法，在MovingAI基准数据集上验证了'
     '加权方法的加速效果，但其权重为预设常数，在搜索过程中保持不变，且加权后路径代价显著增加。冯志乾等[7]将量化的'
     '障碍物信息作为启发函数的自适应调节权重，提高了算法搜索效率，但其权重基于起点到终点矩形区域的全局障碍率，'
     '在搜索过程中为固定值，无法根据局部环境动态调整。高建京等[8]引入启发函数动态加权方法，根据当前节点到目标点'
     '的距离比例调整权重，减少了寻路时间，但该方法的权重仅依赖启发值本身的衰减，未利用环境结构信息。毕竟等[9]'
     '引入障碍率概念设计评价函数f(n)=g(n)+(1−lnP)×h(n)，但其障碍率基于起点与目标点间的全局统计，在搜索过程中为'
     '固定值。')
para('在路径平滑与优化方面，赵江等[10]借助几何方法对A*算法规划的路径进行优化，缩短了运输路径长度和减少了转折次数，'
     '但路径仍存在不平滑转折点。冯泽鹏等[11]提出三角形边界与三邻域搜索相结合的搜索机制，引入三次B样条曲线平滑'
     '处理，在搜索时间和遍历节点方面取得显著改进，但未考虑局部障碍物密度对搜索策略的影响。赵晓等[12]采用跳点搜索'
     '算法改进A*，搜索节点明显减少，但无法实现动态避障。吴鹏等[13]提出双向搜索及自适应权重算法，减少了路径规划'
     '时间，但路径平滑性不足。')
para('综上所述，现有改进A*算法在启发函数权重设计方面普遍采用全局固定或仅基于距离衰减的策略，未能充分利用局部'
     '环境结构信息实现逐节点的搜索策略自适应调整。此外，多数研究采用自设计地图进行验证，缺乏在公开基准数据集上'
     '的系统性对比和统计检验。')
para('针对上述不足，本文提出残差驱动自适应加权A*路径规划算法（RDA*），主要贡献如下：')
para('（1）提出基于局部障碍物密度的逐节点自适应权重机制。通过积分图实现O(1)复杂度的局部障碍率查询，定义通行残差'
     'r_res(n)=1−ρ_local(n)作为驱动因子，设计权重函数α(n)=1+β·r_res(n)，使算法在空旷区域增大启发权重加速搜索，'
     '在障碍物密集区域降低权重保守寻优，实现搜索效率与路径质量的自适应平衡。')
para('（2）设计两阶段路径平滑策略。第一阶段基于严格supercover视线检测的路径简化，第二阶段基于角点加权插值的路径'
     '平滑，在保证路径合法性的前提下显著缩短路径长度并减少转弯次数。')
para('（3）在MovingAI公开基准数据集的15张地图上，采用7种算法（含2组消融）×20条任务的大规模实验设计，结合Wilcoxon'
     '符号秩检验和BH-FDR校正进行统计分析，系统验证了算法的有效性和鲁棒性。')

# ═════════════ 1 算法设计 ═════════════
heading('1　算法设计', 1)
heading('1.1　传统A*算法', 2)
para('A*算法是一种基于图搜索的启发式路径规划算法[5]，其代价函数定义为：')
formula('f(n) = g(n) + h(n)', '（1）')
para('其中，n为当前节点；g(n)为从起点到节点n的实际代价；h(n)为从节点n到目标点的估计代价（启发函数）。A*算法通过'
     '优先扩展f(n)值最小的节点，在启发函数满足可接受性h(n)≤h*(n)的条件下，能够保证找到最优路径。在八连通栅格地图'
     '中，欧氏距离虽然满足可接受性，但会低估对角移动的代价，导致扩展大量不必要的节点。')

heading('1.2　Octile距离启发函数', 2)
para('在八连通栅格地图中，Octile距离是最紧凑的可接受启发函数，其定义为：')
formula('h_oct(n) = max(dx, dy) + (√2 − 1) · min(dx, dy)', '（2）')
para('其中，dx=|x_n−x_g|，dy=|y_n−y_g|分别为当前节点与目标点在x、y方向上的距离差。Octile距离精确描述了八连通网格'
     '中的最短移动距离，相比欧氏距离能更准确地估计实际路径代价，从而减少不必要的节点扩展。')

heading('1.3　残差驱动自适应权重机制', 2)

# 1.3.1 局部障碍物密度 (含文献来源)
heading('1.3.1　局部障碍物密度', 3)
para('在启发函数权重设计中，障碍物密度已被用于评估路径搜索的环境复杂度。冯志乾等[7]将起点与目标点组成的矩形区域'
     '内的障碍物面积占比作为权重因子，毕竟等[9]将起终点间的全局障碍率纳入评价函数。上述方法的障碍物密度均基于'
     '全局统计，在搜索过程中为固定值，无法反映不同节点处的局部环境差异。')
para('为实现逐节点的搜索策略自适应调整，本文将障碍物密度的计算范围从全局缩小到节点邻域。在栅格地图中，障碍物'
     '栅格取值为1，自由栅格取值为0，则任意区域内栅格值的算术均值即为该区域的障碍物占比。据此，定义局部障碍物密度'
     'ρ_local(n)为以节点n为中心、边长为(2r+1)的正方形窗口W_r(n)内所有栅格值的均值：')
formula('ρ_local(n) = Σ grid(i,j) / (2r+1)²', '（3）')
para('其中，grid(i,j)∈{0,1}为栅格地图的占用值，0表示自由栅格，1表示障碍物栅格；(2r+1)²为窗口内的栅格总数。由定义'
     '可知，ρ_local(n)∈[0,1]：当窗口内全为自由栅格时ρ_local(n)=0，当窗口内全为障碍物时ρ_local(n)=1。本文取r=5，'
     '窗口大小为11×11，该参数的选择论证见1.3.4节。')
para('若对每个扩展节点直接遍历其(2r+1)×(2r+1)窗口统计障碍物数量，单次查询的时间复杂度为O(k²)（k=2r+1=11），在搜索'
     '过程中需执行数万次查询，将显著增加计算开销，削弱自适应权重带来的加速收益。为此，本文引入积分图（integral '
     'image）技术实现O(1)复杂度的局部障碍率查询。具体而言，在搜索前对障碍物栅格预计算二维前缀和矩阵I(x,y)=Σgrid'
     '(i,j)（i≤x, j≤y），时间复杂度为O(H×W)，仅执行一次。查询时，任意矩形区域(x₁,y₁)至(x₂,y₂)内的障碍物总数可通过'
     'I(x₂,y₂)−I(x₁−1,y₂)−I(x₂,y₁−1)+I(x₁−1,y₁−1)在常数时间内获得。该方法确保了逐节点障碍率查询不成为搜索瓶颈，'
     '为残差驱动的自适应权重调节提供了实时性保障。')

# 1.3.2 通行残差 (r_res符号)
heading('1.3.2　通行残差定义与残差驱动机制', 3)
para('在路径搜索过程中，节点周围的障碍物分布决定了该区域的通行难度。ρ_local(n)刻画了节点n周围的障碍物"填充程度"，'
     '但搜索策略的调整应当依据该节点的"通行裕度"——即当前环境偏离完全阻塞状态的程度。为此，定义通行残差r_res(n)：')
formula('r_res(n) = 1 − ρ_local(n)', '（4）')
para('其中，ρ_local(n)为式（3）定义的局部障碍物密度。r_res(n)∈[0, 1]，其物理意义为：节点n的11×11邻域窗口内，自由'
     '栅格占总栅格的比例，即局部环境的可通行空间残余率。')
para('通行残差r_res(n)具有以下边界性质：（1）当r_res(n)→1时，节点周围几乎无障碍物，通行裕度充足，搜索空间开阔，'
     '此时存在大量等价或近似等价的路径，标准A*算法会扩展大量不必要的节点；（2）当r_res(n)→0时，节点周围障碍物密集，'
     '通行裕度极低，可行路径稀少且路径间差异显著，此时搜索需要精细评估每条路径的代价以保证质量。')

# 1.3.3 自适应权重函数 (r_res符号)
heading('1.3.3　自适应权重函数', 3)
para('基于通行残差r_res(n)，构造自适应权重函数：')
formula('α(n) = 1 + β · r_res(n)', '（5）')
para('将式（4）代入式（5），得到与局部障碍物密度的显式关系：')
formula('α(n) = 1 + β · (1 − ρ_local(n))', '（6）')
para('其中β=0.3，α(n)∈[1.0, 1.3]。改进后的代价函数为：')
formula('f(n) = g(n) + α(n) · h_oct(n) = g(n) + [1 + β · r_res(n)] · h_oct(n)', '（7）')
para('式（7）可展开为：')
formula('f(n) = [g(n) + h_oct(n)] + β · r_res(n) · h_oct(n)', '（8）')
para('其中，g(n)+h_oct(n)为标准A*（Octile启发）的代价函数，β·r_res(n)·h_oct(n)为残差驱动的加速项。该加速项的大小'
     '同时取决于两个因素：通行残差r_res(n)反映局部环境的通行裕度，h_oct(n)反映当前节点到目标点的剩余距离。因此，'
     '残差驱动机制的调控逻辑为：仅在通行裕度充足且距目标较远时施加显著的搜索加速，而在通行裕度不足或已接近目标时，'
     '加速项自然衰减，算法退化为标准A*以保证路径质量。')

# 1.3.4 参数选择 (含r敏感性)
heading('1.3.4　参数选择', 3)
para('参数β通过在MovingAI基准数据集的调参集上（15张地图，30%−50%分位任务，与评估集零重叠）进行网格搜索确定。'
     '以"扩展节点减少率最大且路径长度增加率<5%"为约束，最终选定β*=0.3，对应的α(n)∈[1.0, 1.3]。')
para('窗口半径r决定了局部障碍物密度ρ_local(n)的感知范围。r过小时，窗口内栅格数少，密度估计对单个障碍物敏感，容易'
     '产生噪声波动，导致权重α(n)在相邻节点间剧烈变化；r过大时，窗口覆盖范围趋近全局，密度估计被过度平滑，不同区域'
     '的ρ_local(n)差异缩小，自适应调节能力减弱。为确定r的最优取值，在调参集上固定β=0.3，分别测试r∈{3, 5, 7}三种'
     '窗口大小，以A*(Octile)为基线，统计15张地图上扩展节点减少率和路径长度增加率的均值，结果如表1所示。')
caption('表1　不同窗口半径r的实验结果', 'Table 1　Experimental results of different window radii r')
tbl(['r','窗口大小','栅格数','平均EXP减少率/%','平均PL增加率/%','满足PL<5%'],
    [['3','7×7','49','54.50','4.07','是'],['5','11×11','121','53.46','3.80','是'],['7','15×15','225','52.42','3.42','是']])
para('由表1可知，三种r取值的性能差异较小（EXP减少率在52%−55%之间，PL增加率在3.4%−4.1%之间），均满足PL<5%的约束。'
     '综合考虑，本文选取r=5（11×11窗口）：相比r=3，密度估计更稳定，避免因个别障碍物导致的权重剧烈波动；相比r=7，'
     '保持了更高的空间分辨率，能够更敏锐地区分空旷区域与密集区域的边界。')

# 1.4 两阶段路径平滑
heading('1.4　两阶段路径平滑', 2)
heading('1.4.1　第一阶段：视线简化', 3)
para('视线简化的核心思想是：若路径上两个不相邻的节点之间存在无障碍视线，则中间节点可被移除。从路径起点开始，'
     '保留当前节点为锚点；依次检测锚点与后续节点之间的视线（line-of-sight）；当视线被障碍物阻断时，保留阻断前的'
     '最后一个可见节点并设为新锚点。视线检测采用严格的supercover算法，确保路径不会穿越任何障碍物。')
heading('1.4.2　第二阶段：角点插值平滑', 3)
para('对于简化后的路径，仍可能存在锐角转折。角点插值通过加权平均将转折点向平滑方向偏移。对于路径中的每个非端点'
     '节点B（前驱为A，后继为C），计算平滑点M = ((A+2B+C)/4)。平滑点M仅在不位于障碍物、且与前后节点视线均无障碍'
     '时替代原节点。此设计确保平滑过程中使用的是已平滑的前驱节点，避免连续平滑导致的路径非法问题。')

# 1.5 理论性质分析 (完整证明)
heading('1.5　算法理论性质分析', 2)
heading('1.5.1　收敛性证明', 3)
para('命题1：在有限大小的静态栅格地图上，若起点s与目标点goal之间存在可行路径，则RDA*算法一定在有限步内终止并返回'
     '一条从s到goal的可行路径。')
para('证明：RDA*算法的搜索过程维护开放列表（open list）和关闭列表（closed list）。每次迭代从开放列表中取出f值最小'
     '的节点进行扩展，扩展后将其移入关闭列表。对该过程的终止性和可行性分析如下：')
para('（1）有限性。设栅格地图大小为H×W，则可扩展的自由栅格节点总数N_free≤H×W为有限值。每个节点至多被扩展一次'
     '（关闭列表保证已扩展的节点不再被重复取出），因此搜索过程至多执行N_free次迭代。')
para('（2）代价正定性。八连通栅格中任意相邻节点间的步进代价c(n,n\')≥1>0，因此从起点到任意节点的实际代价g(n)严格'
     '递增。结合α(n)≥1和h_oct(n)≥0，代价函数f(n)=g(n)+α(n)·h_oct(n)>0对所有非起点节点成立，不存在零代价循环。')
para('（3）可达性。因为起点s与目标点goal之间存在可行路径，且八连通邻域搜索能够覆盖所有相邻的自由栅格节点，因此'
     '目标点goal必然会被发现并加入开放列表。开放列表非空时搜索持续进行，goal最终会被取出，搜索终止。')
para('（4）路径合法性。搜索过程中的邻域扩展函数neighbors8仅返回自由栅格节点，且对角移动采用AND逻辑约束（即对角方向'
     '的两个正交邻居均为自由栅格时才允许移动），因此搜索得到的路径不会穿越障碍物。')
para('综合（1）−（4），RDA*算法在有限步内终止并返回合法的可行路径。证毕。')

heading('1.5.2　路径代价上界', 3)
para('命题2：设最优路径代价为C*，RDA*算法在搜索阶段（平滑前）返回的路径代价为C_search，则C_search≤(1+β)·C*。')
para('证明：RDA*算法的代价函数为f(n)=g(n)+α(n)·h_oct(n)，其中α(n)∈[1, 1+β]。令w_max=1+β=1.3，则对任意节点n有：')
formula('α(n) · h_oct(n) ≤ w_max · h_oct(n)', '（9）')
para('因此，RDA*的代价函数满足：')
formula('f_RDA(n) ≤ g(n) + w_max · h_oct(n) = f_WA(n)', '（10）')
para('其中f_WA(n)为使用固定权重w_max的标准加权A*的代价函数。根据加权A*算法的经典理论[5]，当启发函数h(n)满足'
     '可接受性时，固定权重w_max的WA*算法返回的路径代价C_WA满足：')
formula('C_WA ≤ w_max · C*', '（11）')
para('当RDA*终止时，目标节点goal从开放列表中取出，h_oct(goal)=0，因此C_search=g(goal)。由于RDA*对每个节点的f值'
     '不超过WA*对同一节点的f值（式（10）），结合有限图上加权A*的有界次优性保证，RDA*搜索阶段返回的路径代价满足：')
formula('C_search ≤ (1+β) · C* = 1.3 · C*', '（12）')
para('当β=0.3时，理论路径代价上界为1.3·C*，即搜索阶段的路径长度最多比最优路径增加30%。需要指出的是，式（12）为'
     '搜索阶段的理论上界，经过两阶段路径平滑后，实际路径代价将进一步降低。实验结果表明，RDA*最终输出的路径长度仅'
     '比A*(Octile)的最优路径增加3.8%，远优于30%的理论上界，说明两阶段平滑策略有效补偿了加权搜索带来的路径次优性。')

# 1.6 时间复杂度
heading('1.6　时间复杂度分析', 2)
para('设地图大小为H×W，路径节点数为n。RDA*算法的各阶段复杂度如下：积分图构建O(H×W)，仅执行一次；逐节点障碍率'
     '查询O(1)；A*搜索O(N\'·logN\')，其中N\'为改进后的扩展节点数，显著小于传统A*的N；路径简化O(n·L)，其中L为视线'
     '检测的平均扫描长度；角点平滑O(n)。整体时间复杂度与传统A*同阶，但因扩展节点大幅减少而实际运行更快。')

# ═════════════ 2 实验设计 ═════════════
heading('2　实验设计', 1)
heading('2.1　实验环境与数据集', 2)
para('实验在Ubuntu 24.04操作系统下进行，编程语言为Python 3.12。本文采用MovingAI公开基准数据集[4]进行实验验证。'
     '从数据集中选取15张地图，覆盖三种类型和五个障碍率梯度，具体如表2所示。')
caption('表2　实验地图信息', 'Table 2　Information of experimental maps')
tbl(['类型','地图名称','尺寸','障碍率/%'],
    [['DAO','brc503d','257×320','18.1'],['DAO','orz100d','395×412','38.8'],
     ['DAO','den502d','251×211','48.6'],['DAO','arena2','209×281','58.6'],
     ['DAO','ost003d','194×194','64.9'],
     ['Street','Berlin_0_512','512×512','25.0'],['Street','London_0_512','512×512','25.0'],
     ['Street','Moscow_0_512','512×512','25.0'],['Street','Paris_0_512','512×512','25.0'],
     ['Street','Shanghai_0_512','512×512','25.0'],
     ['WC3','dustwallowkeys','512×512','31.5'],['WC3','blastedlands','512×512','49.9'],
     ['WC3','bloodvenomfalls','512×512','55.7'],['WC3','darkforest','512×512','61.9'],
     ['WC3','battleground','512×512','64.8']])

heading('2.2　对比算法', 2)
para('实验设置7种算法进行对比，包括4种基线算法、本文算法及2种消融变体，具体如表3所示。')
caption('表3　对比算法设置', 'Table 3　Configuration of compared algorithms')
tbl(['编号','算法名称','启发函数','权重','平滑'],
    [['G1','Dijkstra','h=0','—','无'],['G2','A*(欧氏)','欧氏距离','1.0','无'],
     ['G3','A*(Octile)','Octile距离','1.0','无'],['G4','WA*(1.2)','Octile距离','1.2固定','无'],
     ['G5','RDA*（本文）','Octile距离','α(n)自适应','两阶段'],
     ['G6','NoAdaptive','Octile距离','1.0','两阶段'],
     ['G7','NoSmooth','Octile距离','α(n)自适应','无']])

heading('2.3　评价指标与统计方法', 2)
para('采用路径长度（PL）、转弯次数（TC）、扩展节点数（EXP）、搜索时间（search_ms）、总时间（total_ms）和碰撞检测（CF）'
     '6项指标评价算法性能。对每张地图取20条任务的均值作为该地图的代表值，得到n=15个样本对。采用Wilcoxon符号秩'
     '检验（双侧）进行配对比较，使用BH方法进行FDR校正。同时报告Cohen\'s d效应量和Win/Tie/Loss统计。')

# ═════════════ 3 实验结果 ═════════════
heading('3　实验结果与分析', 1)

heading('3.1　参数选择', 2)
para('在30%−50%分位的调参集上，对β∈{0.1, 0.2, 0.3, 0.4, 0.5}进行网格搜索，结果如表4所示。β的取值面临搜索效率与'
     '路径质量之间的权衡：β越大，启发权重越高，扩展节点减少越多，搜索速度越快，但路径偏离最优的程度也越大。为此，'
     '设定路径长度增加率<5%作为约束条件（超过5%的路径质量损失在工程应用中通常不可接受）。由表4可知，β=0.4和β=0.5'
     '的PL增加率分别为5.7%和7.1%，超出约束范围；β=0.1和β=0.2虽然路径质量损失更小，但搜索加速效果有限（EXP减少率'
     '分别为35.5%和46.4%）；β=0.3在满足PL增加率<5%约束的前提下，实现了最大的EXP减少率（52.2%），为搜索效率与路径'
     '质量的最优平衡点，因此选定β*=0.3。')
caption('表4　β参数调优结果', 'Table 4　Results of β parameter tuning')
tbl(['β','平均EXP减少率/%','平均PL增加率/%','是否可行'],
    [['0.1','35.5','1.4','是'],['0.2','46.4','2.9','是'],
     ['0.3','52.2','4.5','是（选定）'],['0.4','56.0','5.7','否（PL>5%）'],
     ['0.5','59.5','7.1','否（PL>5%）']])
img('figures/chart_beta_sensitivity.png', 12)
caption('图1　β参数敏感性分析', 'Fig.1　β parameter sensitivity analysis')

para('为验证β=0.3在不同障碍率场景下的适用性，将15张地图按障碍率分为三组：低障碍率（<30%，6张）、中障碍率（30%−55%，'
     '4张）、高障碍率（>55%，5张），分别统计各组在β=0.3下的性能表现，结果如表5所示。')
caption('表5　β=0.3在不同障碍率分组下的性能表现', 'Table 5　Performance of β=0.3 across different obstacle ratio groups')
tbl(['障碍率分组','地图数','平均EXP减少率/%','平均PL变化率/%','满足PL<5%'],
    [['低(<30%)','6','65.8','−3.65','是'],
     ['中(30%-55%)','4','49.8','−3.89','是'],
     ['高(>55%)','5','41.6','−4.26','是']])
para('由表5可知，β=0.3在三种障碍率场景下均满足PL<5%的约束条件，且性能表现呈现出与算法设计一致的趋势：在低障碍率'
     '地图中，局部通行残差普遍较大，自适应权重增益显著，扩展节点减少率最高（65.8%）；在高障碍率地图中，通行残差普遍'
     '较小，权重自动趋近1.0，扩展节点减少率降低（41.6%），但路径质量仍然达标。这一结果表明，β=0.3无需针对不同障碍率'
     '场景单独调整，自适应权重机制本身已通过通行残差实现了场景适配，验证了该参数选择的鲁棒性。')

heading('3.2　主实验结果', 2)
para('主实验在15张地图上以20条任务×7种算法执行，共计2100次搜索。所有搜索均成功完成，超时率为0，碰撞合法率为100%。'
     '全局均值结果如表6所示。')
caption('表6　主实验全局均值结果', 'Table 6　Global mean results of main experiment')
tbl(['算法','PL','TC','EXP','search_ms','total_ms'],
    [['Dijkstra','312.94','19.6','81568','559.21','559.21'],
     ['A*(欧氏)','312.94','51.1','17019','126.31','126.31'],
     ['A*(Octile)','312.94','22.5','13331','113.43','113.43'],
     ['WA*(1.2)','318.53','37.5','7171','60.49','60.49'],
     ['RDA*（本文）','301.76','12.1','6550','80.67','86.04'],
     ['NoAdaptive','299.53','8.7','13331','113.22','119.08'],
     ['NoSmooth','328.41','57.9','6550','80.52','80.52']])
para('由表6可知，与A*(欧氏)相比，本文RDA*算法在所有核心指标上均取得显著改善：路径长度缩短3.6%，转弯次数减少76.4%，'
     '扩展节点减少61.5%，总运行时间缩短31.9%。')
img('figures/chart_expanded_nodes.png', 14)
caption('图2　15张地图扩展节点数对比', 'Fig.2　Comparison of expanded nodes across 15 maps')
img('figures/chart_path_length.png', 14)
caption('图3　15张地图路径长度对比', 'Fig.3　Comparison of path length across 15 maps')
img('figures/chart_turn_count.png', 14)
caption('图4　15张地图转弯次数对比', 'Fig.4　Comparison of turn count across 15 maps')
img('figures/expanded_nodes_visualization.png', 16)
caption('图5　扩展节点可视化对比（bloodvenomfalls, obs=55.7%）', 
        'Fig.5　Visualization of expanded nodes (bloodvenomfalls, obs=55.7%)')
para('图5直观展示了自适应权重机制对搜索范围的缩减效果。A*(Octile)和NoAdaptive（消融）的扩展节点区域完全相同'
     '（EXP=12905），而RDA*的扩展区域大幅缩小（EXP=2057，减少84%），搜索聚焦于路径周围的狭窄带状区域。NoAdaptive'
     '与A*(Octile)的搜索范围一致，验证了扩展节点的减少完全来自自适应权重模块的贡献，而非路径平滑的间接影响。')

# 3.3 统计检验 (含A*(Octile)补充段)
heading('3.3　统计检验', 2)
para('对RDA*与A*(欧氏)的配对比较进行Wilcoxon符号秩检验，结果如表7所示。')
caption('表7　统计检验结果', 'Table 7　Statistical test results of RDA* vs A*(Euclidean)')
tbl(['指标','变化率/%','p_raw','q_bh','显著性','W/T/L'],
    [['PL','−3.6','0.000061','0.000105','***','15/0/0'],
     ['TC','−76.4','0.000061','0.000105','***','15/0/0'],
     ['EXP','−61.5','0.000061','0.000105','***','15/0/0'],
     ['total_ms','−31.9','0.015','0.020','*','12/0/3'],
     ['search_ms','−36.1','0.012','0.018','*','12/0/3']])
para('PL、TC、EXP三项核心指标的q_bh值均远小于0.001（极显著），W/T/L=15/0/0，即RDA*在所有15张地图上均优于A*(欧氏)。')

# ★ 新增A*(Octile)对比段
para('需要指出的是，A*(Octile)作为采用最紧凑可接受启发函数的基线，其路径长度与A*(欧氏)一致（均为312.94），但扩展'
     '节点更少（13331 vs 17019）。由表6可知，RDA*相对A*(Octile)仍取得显著改善：扩展节点减少50.9%（6550 vs 13331），'
     '转弯次数减少46.2%（12.1 vs 22.5），路径长度缩短3.6%（301.76 vs 312.94），总运行时间缩短24.2%（86.04 vs '
     '113.43）。这表明RDA*的性能提升并非仅来自启发函数从欧氏距离到Octile距离的替换，自适应权重机制和两阶段平滑在'
     'Octile基线之上仍提供了实质性的增益。')


# 3.4 消融 (表8修正 == → —，加表注)
heading('3.4　消融分析', 2)
para('为验证自适应权重和路径平滑两个模块各自的贡献，对RDA*与两种消融变体进行对比分析，结果如表8所示。')
caption('表8　消融实验结果', 'Table 8　Results of ablation study')
tbl(['对比','指标','变化率/%','p_raw','显著性'],
    [['RDA* vs NoAdaptive','EXP','−50.9','0.000061','***'],
     ['RDA* vs NoAdaptive','total_ms','−27.7','0.030','*'],
     ['RDA* vs NoAdaptive','PL','+0.7','0.000061','***'],
     ['RDA* vs NoSmooth','PL','−8.1','0.000061','***'],
     ['RDA* vs NoSmooth','TC','−79.2','0.000061','***'],
     ['RDA* vs NoSmooth','EXP','0','—','—']],
    footnote='注：显著性标记中，***表示p<0.001，*表示p<0.05，—表示两组数据完全相同，无统计差异。')
para('消融结果表明：（1）自适应权重模块的主要贡献在于搜索效率——使扩展节点减少50.9%，代价是路径长度微增0.7%；'
     '（2）路径平滑模块的主要贡献在于路径质量——使路径长度缩短8.1%，转弯次数减少79.2%。两个模块功能正交、互为补充。')
img('figures/chart_ablation.png', 12)
caption('图6　消融实验对比', 'Fig.6　Ablation study comparison')

# 3.5 困难场景
heading('3.5　困难场景验证', 2)
para('为验证算法在困难场景下的鲁棒性，在相同15张地图上使用最长路径任务（平均最优路径≈626步）进行补充实验。结果'
     '如表9所示。')
caption('表9　主实验与补充实验对比', 'Table 9　Comparison of main and supplementary experiments')
tbl(['指标','主实验(PL≈316)','补充实验(PL≈626)'],
    [['PL变化率','−3.6% ***','−2.8% ***'],
     ['TC变化率','−76.4% ***','−71.9% ***'],
     ['EXP变化率','−61.5% ***','−58.4% ***']])
para('补充实验的改进趋势与主实验完全一致，三项核心指标的p值均小于0.001，说明RDA*在困难场景下仍具有良好的鲁棒性。')

# ═════════════ 4 结论 (r_res符号) ═════════════
heading('4　结论', 1)
para('本文针对传统A*算法在复杂栅格环境中搜索效率受限、路径冗余转折多、启发函数权重无法自适应调整等问题，提出了'
     '残差驱动自适应加权A*路径规划算法（RDA*），通过融合局部障碍物密度感知与两阶段路径平滑策略，实现了搜索效率与'
     '路径质量的协同优化。')
para('在搜索策略方面，本文以Octile距离替代传统欧氏距离作为启发函数，更准确地估计八连通栅格中的实际路径代价。在此'
     '基础上，引入局部障碍物密度概念，利用积分图实现O(1)复杂度的逐节点障碍率查询，定义通行残差r_res(n)=1−ρ_local(n)'
     '作为权重调节的驱动因子，设计自适应权重函数α(n)=1+β·r_res(n)，使算法能够根据节点周围的环境特征动态调整搜索'
     '策略——在空旷区域增大启发权重以加速搜索，在障碍物密集区域降低权重以保守寻优。通过在调参集上的网格搜索确定'
     'β*=0.3，在路径长度增加率控制在5%以内的前提下，实现了52.2%的扩展节点减少率。')
para('在路径优化方面，本文提出两阶段平滑策略：第一阶段基于严格supercover视线检测进行路径简化，去除冗余中间节点；'
     '第二阶段通过角点加权插值平滑路径转折，降低路径曲率变化。两个阶段均内置碰撞合法性检查，确保平滑后的路径不'
     '穿越任何障碍物。该策略有效补偿了加权搜索带来的路径次优性，使最终路径长度反而优于标准A*算法。')
para('实验验证方面，本文在MovingAI公开基准数据集的15张地图上进行了大规模系统性实验。主实验以7种算法×20条任务共计'
     '2100次搜索为基础，辅以Wilcoxon符号秩检验和BH-FDR校正进行统计分析。结果表明，与A*(欧氏)相比，RDA*的路径长度'
     '缩短3.6%，转弯次数减少76.4%，扩展节点减少61.5%，总运行时间缩短31.9%，四项指标均达到统计显著水平。消融实验'
     '证实了自适应权重模块与路径平滑模块各自具有独立且显著的贡献：前者使扩展节点减半，后者使路径长度缩短8.1%、'
     '转弯次数减少79.2%，两个模块功能正交、互为补充。补充实验在长路径场景（平均最优路径≈626步）下进一步验证了算法'
     '的鲁棒性，改进趋势与主实验完全一致。')
para('同时，本文也存在以下局限：（1）在障碍物分布均匀的环境中，局部障碍率的空间差异较小，自适应权重退化为近似固定值，'
     '算法优势减弱；（2）当前实验仅在仿真环境中进行，尚未在实际机器人平台上验证算法的实时性和可行性。')
para('未来的研究可从以下方向展开：引入β参数的自动衰减机制，使算法在障碍物分布均匀的环境中也能保持竞争力；增加JPS、'
     'Theta*等高级路径规划算法作为对比基线，进一步验证算法的相对优势；开展Gazebo仿真与真实机器人平台的实验，验证'
     '算法在动态环境中的实际应用效果。')

# ═════════════ 参考文献 (按首次引用顺序+待发表格式) ═════════════
heading('参考文献：', 1)
refs = [
    '[1] 朱大奇, 颜明重. 移动机器人路径规划技术综述[J]. 控制与决策, 2010, 25(7): 961-967.',
    '[2] 王旭, 朱其新, 朱永红. 面向二维移动机器人的路径规划算法综述[J]. 计算机工程与应用, 2023, 59(20): 51-66.',
    '[3] 毛建旭, 贺振宇, 王耀南, 等. 电力巡检机器人路径规划技术及应用综述[J]. 控制与决策, 2023, 38(11): 3009-3024.',
    '[4] STURTEVANT N R. Benchmarks for grid-based pathfinding[J]. IEEE Transactions on Computational Intelligence and AI in Games, 2012, 4(2): 144-148.',
    '[5] HART P E, NILSSON N J, RAPHAEL B. A formal basis for the heuristic determination of minimum cost paths[J]. IEEE Transactions on Systems Science and Cybernetics, 1968, 4(2): 100-107.',
    '[6] ALGFOOR Z A, SUNAR M S, ABDULLAH A. A new weighted pathfinding algorithms to reduce the search time on grid maps[J]. Expert Systems with Applications, 2017, 71: 319-331.',
    '[7] 冯志乾, 王欣, 吴迪. 基于改进A*和DWA融合的移动机器人路径规划[J]. 计算机应用与软件, 2026, 43(2): 340-346.',
    '[8] 高建京, 崔明月. 融合改进A*算法和动态窗口法的移动机器人路径规划研究[J]. 控制工程, 待发表.',
    '[9] 毕竟, 刘俊. 结合DC-A*与FE-DWA的巡检机器人路径规划方法[J]. 计算机工程与应用, 2026, 62(3): 334-346.',
    '[10] 赵江, 张岩, 马泽文. AGV路径规划A星算法的改进与验证[J]. 计算机工程与应用, 2018, 54(21): 217-223.',
    '[11] 冯泽鹏, 李宗刚, 夏广庆, 等. 改进A*与APF的移动机器人路径规划算法研究[J]. 计算机工程与应用, 2025, 61(20): 132-145.',
    '[12] 赵晓, 王铮, 黄程侃. 基于改进A*算法的移动机器人路径规划[J]. 机器人, 2018, 40(6): 903-910.',
    '[13] 吴鹏, 桑成军, 陆忠华, 等. 基于改进A*算法的移动机器人路径规划研究[J]. 计算机工程与应用, 2019, 55(21): 227-233.',
]
for ref in refs:
    para(ref, sz=9, indent=False)

out = '/mnt/user-data/outputs/RDA_star_论文修改稿_v4.docx'
doc.save(out)
print(f"✅ {out}")
