#!/usr/bin/env python3
"""生成 RDA* 论文完整 Word 初稿"""
import os
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# ── 页面设置 (A4) ──
for section in doc.sections:
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)

style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.paragraph_format.line_spacing = 1.25
style.paragraph_format.space_after = Pt(0)

def add_heading_cn(text, level=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if level == 0:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    if level == 0:
        run.font.size = Pt(16)
    elif level == 1:
        run.font.size = Pt(14)
    elif level == 2:
        run.font.size = Pt(12)
    elif level == 3:
        run.font.size = Pt(11)
    run.font.name = '黑体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    return p

def add_para(text, bold=False, size=10.5, font='宋体', align=None, indent=True):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = font
    run.element.rPr.rFonts.set(qn('w:eastAsia'), font)
    run.bold = bold
    return p

def add_table_caption(cn_text, en_text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(cn_text)
    run.bold = True
    run.font.size = Pt(9)
    run.font.name = '黑体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(en_text)
    run2.bold = True
    run2.font.size = Pt(9)
    run2.font.name = 'Times New Roman'

def add_fig_caption(cn_text, en_text):
    add_table_caption(cn_text, en_text)

def add_simple_table(headers, rows):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.name = 'Times New Roman'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri+1].cells[ci]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            run.font.name = 'Times New Roman'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            if 'RDA*' in str(val) or '本文' in str(val):
                run.bold = True
    return table

def add_image(path, width_cm=14):
    if os.path.exists(path):
        doc.add_picture(path, width=Cm(width_cm))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ════════════════════════════════════════
# 标题
# ════════════════════════════════════════
add_heading_cn('残差驱动自适应加权A*路径规划算法', level=0)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('刘家树')
run.font.size = Pt(12)
run.font.name = '楷体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('南方科技大学 机械与能源工程系，广东 深圳 518055')
run.font.size = Pt(9)
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# ════════════════════════════════════════
# 中文摘要
# ════════════════════════════════════════
p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(0.74)
run = p.add_run('摘　要：')
run.bold = True
run.font.size = Pt(10.5)
run.font.name = '黑体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
run2 = p.add_run(
    '针对A*算法在复杂栅格环境中搜索效率受限、路径质量不佳的问题，提出了一种残差驱动自适应加权A*路径规划算法'
    '（Residual-Driven Adaptive Weighted A*，RDA*）。该算法以Octile距离作为启发函数，引入局部障碍物密度概念，'
    '通过积分图实现O(1)复杂度的逐节点障碍率查询，设计自适应权重函数α(n)=1+β(1−ρ_local(n))，使搜索在空旷区域'
    '加速探索、在密集区域保守寻优；同时采用两阶段路径平滑策略（视线简化与角点插值），在保证路径合法性的前提下'
    '显著缩短路径长度并减少转弯次数。在MovingAI公开基准数据集的15张地图上，以20条任务×7种算法进行对比实验，'
    '结果表明：与A*(欧氏)相比，RDA*的路径长度缩短3.6%（p<0.001），转弯次数减少76.4%（p<0.001），扩展节点减少'
    '61.5%（p<0.001），总运行时间缩短31.9%；消融实验验证了自适应权重与路径平滑两个模块各自的独立贡献。补充实验'
    '在长路径场景下进一步验证了算法的鲁棒性。'
)
run2.font.size = Pt(10.5)
run2.font.name = '宋体'
run2.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(0.74)
run = p.add_run('关键词：')
run.bold = True
run.font.size = Pt(10.5)
run.font.name = '黑体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
run2 = p.add_run('路径规划；A*算法；自适应权重；障碍物密度；路径平滑；栅格地图')
run2.font.size = Pt(10.5)

p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(0.74)
run = p.add_run('文献标志码：')
run.bold = True
run.font.size = Pt(10.5)
run2 = p.add_run('A　中图分类号：TP242')
run2.font.size = Pt(10.5)

# ════════════════════════════════════════
# 英文标题和摘要
# ════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(12)
run = p.add_run('Residual-Driven Adaptive Weighted A* Path Planning Algorithm')
run.bold = True
run.font.size = Pt(12)
run.font.name = 'Times New Roman'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('LIU Jiashu')
run.font.size = Pt(10)
run.font.name = 'Times New Roman'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Department of Mechanical and Energy Engineering, Southern University of Science and Technology, Shenzhen, Guangdong 518055, China')
run.font.size = Pt(9)
run.font.name = 'Times New Roman'

p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(0.74)
run = p.add_run('Abstract: ')
run.bold = True
run.font.size = Pt(9)
run.font.name = 'Times New Roman'
run2 = p.add_run(
    'To address the limited search efficiency and suboptimal path quality of the A* algorithm in complex grid environments, '
    'a Residual-Driven Adaptive Weighted A* (RDA*) path planning algorithm is proposed. The algorithm employs octile distance '
    'as the heuristic function and introduces a local obstacle density concept, achieving O(1) per-node obstacle ratio queries '
    'via integral images. An adaptive weight function α(n)=1+β(1−ρ_local(n)) is designed to accelerate exploration in open areas '
    'while maintaining conservative search in dense regions. A two-stage path smoothing strategy significantly reduces path length '
    'and turn count while ensuring path legality. Experiments on 15 maps from the MovingAI benchmark dataset with 20 tasks × 7 '
    'algorithms show that compared with A*(Euclidean), RDA* reduces path length by 3.6% (p<0.001), turn count by 76.4% (p<0.001), '
    'expanded nodes by 61.5% (p<0.001), and total runtime by 31.9%. Ablation studies confirm the independent contributions of '
    'the adaptive weight and path smoothing modules.'
)
run2.font.size = Pt(9)
run2.font.name = 'Times New Roman'

p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(0.74)
run = p.add_run('Keywords: ')
run.bold = True
run.font.size = Pt(9)
run.font.name = 'Times New Roman'
run2 = p.add_run('path planning; A* algorithm; adaptive weight; obstacle density; path smoothing; grid map')
run2.font.size = Pt(9)
run2.font.name = 'Times New Roman'

# ════════════════════════════════════════
# 0 引言
# ════════════════════════════════════════
add_heading_cn('0　引言', level=1)

add_para(
    '路径规划是实现移动机器人自主导航的关键技术之一，其目标是在已知或部分已知的环境中规划一条从起点到目标点的'
    '最优路径，同时避开障碍物并满足路径长度、转弯次数、搜索效率等多项性能指标[1]。随着机器人技术在物流仓储[2]、'
    '巡检作业[3]、游戏AI[4]等领域的广泛应用，路径规划算法的性能需求日益提高。'
)

add_para(
    'A*算法作为一种经典的启发式搜索算法[5]，通过代价函数f(n)=g(n)+h(n)引导搜索方向，兼顾了Dijkstra算法的最优性'
    '和贪心搜索的高效性，是目前全局路径规划中应用最为广泛的算法之一。然而，传统A*算法在复杂栅格环境中仍存在'
    '以下不足：（1）搜索效率受限，在大规模地图中扩展节点过多，导致运行时间较长；（2）规划路径包含大量冗余转折点，'
    '不利于机器人平稳运行；（3）启发函数权重固定，无法根据局部环境特征自适应调整搜索策略。'
)

add_para(
    '针对上述问题，国内外学者进行了大量研究。在启发函数权重优化方面，Algfoor等[6]提出三种加权启发式搜索技术，'
    '分别基于迭代次数、起终点距离和行进代价设计权重，并应用于A*、双向A*和JPS算法，在MovingAI基准数据集上验证了'
    '加权方法的加速效果，但其权重为预设常数，在搜索过程中保持不变，且加权后路径代价显著增加。冯志乾等[7]将量化的'
    '障碍物信息作为启发函数的自适应调节权重，提高了算法搜索效率，但其权重基于起点到终点矩形区域的全局障碍率，'
    '在搜索过程中为固定值，无法根据局部环境动态调整。高建京等[8]引入启发函数动态加权方法，根据当前节点到目标点'
    '的距离比例调整权重，减少了寻路时间，但该方法的权重仅依赖启发值本身的衰减，未利用环境结构信息。毕竟等[9]'
    '引入障碍率概念设计评价函数f(n)=g(n)+(1−lnP)×h(n)，但其障碍率基于起点与目标点间的全局统计，在搜索过程中为'
    '固定值。'
)

add_para(
    '在路径平滑与优化方面，赵江等[10]借助几何方法对A*算法规划的路径进行优化，缩短了运输路径长度和减少了转折次数，'
    '但路径仍存在不平滑转折点。冯泽鹏等[11]提出三角形边界与三邻域搜索相结合的搜索机制，引入三次B样条曲线平滑'
    '处理，在搜索时间和遍历节点方面取得显著改进，但未考虑局部障碍物密度对搜索策略的影响。赵晓等[12]采用跳点搜索'
    '算法改进A*，搜索节点明显减少，但无法实现动态避障。吴鹏等[13]提出双向搜索及自适应权重算法，减少了路径规划'
    '时间，但路径平滑性不足。'
)

add_para(
    '综上所述，现有改进A*算法在启发函数权重设计方面普遍采用全局固定或仅基于距离衰减的策略，未能充分利用局部'
    '环境结构信息实现逐节点的搜索策略自适应调整。此外，多数研究采用自设计地图进行验证，缺乏在公开基准数据集上'
    '的系统性对比和统计检验。'
)

add_para(
    '针对上述不足，本文提出残差驱动自适应加权A*路径规划算法（RDA*），主要贡献如下：'
)
add_para(
    '（1）提出基于局部障碍物密度的逐节点自适应权重机制。通过积分图实现O(1)复杂度的局部障碍率查询，设计权重函数'
    'α(n)=1+β(1−ρ_local(n))，使算法在空旷区域增大启发权重加速搜索，在障碍物密集区域降低权重保守寻优，实现搜索'
    '效率与路径质量的自适应平衡。'
)
add_para(
    '（2）设计两阶段路径平滑策略。第一阶段基于严格supercover视线检测的路径简化，第二阶段基于角点加权插值的路径'
    '平滑，在保证路径合法性的前提下显著缩短路径长度并减少转弯次数。'
)
add_para(
    '（3）在MovingAI公开基准数据集的15张地图上，采用7种算法（含2组消融）×20条任务的大规模实验设计，结合Wilcoxon'
    '符号秩检验和BH-FDR校正进行统计分析，系统验证了算法的有效性和鲁棒性。'
)

# ════════════════════════════════════════
# 1 算法设计
# ════════════════════════════════════════
add_heading_cn('1　算法设计', level=1)
add_heading_cn('1.1　传统A*算法', level=2)

add_para(
    'A*算法是一种基于图搜索的启发式路径规划算法[5]，其代价函数定义为：'
)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('f(n) = g(n) + h(n)　　（1）')
run.font.name = 'Times New Roman'
run.font.size = Pt(10.5)
run.italic = True

add_para(
    '其中，n为当前节点；g(n)为从起点到节点n的实际代价；h(n)为从节点n到目标点的估计代价（启发函数）。A*算法通过'
    '优先扩展f(n)值最小的节点，在启发函数满足可接受性h(n)≤h*(n)的条件下，能够保证找到最优路径。在八连通栅格地图'
    '中，欧氏距离虽然满足可接受性，但会低估对角移动的代价，导致扩展大量不必要的节点。'
)

add_heading_cn('1.2　Octile距离启发函数', level=2)

add_para(
    '在八连通栅格地图中，Octile距离是最紧凑的可接受启发函数，其定义为：'
)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('h_oct(n) = max(dx, dy) + (√2 − 1) · min(dx, dy)　　（2）')
run.font.name = 'Times New Roman'
run.font.size = Pt(10.5)
run.italic = True

add_para(
    '其中，dx=|x_n−x_g|，dy=|y_n−y_g|分别为当前节点与目标点在x、y方向上的距离差。Octile距离精确描述了八连通网格'
    '中的最短移动距离，相比欧氏距离能更准确地估计实际路径代价，从而减少不必要的节点扩展。'
)

add_heading_cn('1.3　残差驱动自适应权重机制', level=2)
add_heading_cn('1.3.1　局部障碍物密度', level=3)

add_para(
    '为量化节点周围的环境复杂度，定义局部障碍物密度ρ_local(n)为以节点n为中心、半径r的正方形窗口内障碍物栅格的占比：'
)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('ρ_local(n) = Σgrid(i,j) / (2r+1)²　　（3）')
run.font.name = 'Times New Roman'
run.font.size = Pt(10.5)
run.italic = True

add_para(
    '其中，W_r(n)为以节点n为中心、边长为2r+1的正方形窗口；grid(i,j)∈{0,1}，1表示障碍物。本文取r=5，窗口大小为'
    '11×11。为实现高效的局部障碍率查询，预先构建障碍物栅格的积分图（integral image），使任意矩形区域内的障碍物'
    '总数可在O(1)时间内计算。积分图仅需在搜索前构建一次，时间复杂度为O(H×W)。'
)

add_heading_cn('1.3.2　自适应权重函数', level=3)

add_para('基于局部障碍物密度，设计自适应权重函数：')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('α(n) = 1 + β · (1 − ρ_local(n))　　（4）')
run.font.name = 'Times New Roman'
run.font.size = Pt(10.5)
run.italic = True

add_para(
    '其中，β为控制权重调节幅度的参数。该权重函数的物理意义为：当ρ_local(n)较低（空旷区域）时，α(n)趋近1+β，启发'
    '权重增大，搜索加速向目标方向推进；当ρ_local(n)较高（密集区域）时，α(n)趋近1.0，退化为标准A*，保证路径质量。'
    '改进后的代价函数为：'
)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('f(n) = g(n) + α(n) · h_oct(n)　　（5）')
run.font.name = 'Times New Roman'
run.font.size = Pt(10.5)
run.italic = True

add_heading_cn('1.3.3　参数选择', level=3)
add_para(
    '参数β通过在MovingAI基准数据集的调参集上（15张地图，30%−50%分位任务，与评估集零重叠）进行网格搜索确定。'
    '以"扩展节点减少率最大且路径长度增加率<5%"为约束，最终选定β*=0.3，对应的α(n)∈[1.0, 1.3]。'
)

add_heading_cn('1.4　两阶段路径平滑', level=2)
add_heading_cn('1.4.1　第一阶段：视线简化', level=3)
add_para(
    '视线简化的核心思想是：若路径上两个不相邻的节点之间存在无障碍视线，则中间节点可被移除。从路径起点开始，'
    '保留当前节点为锚点；依次检测锚点与后续节点之间的视线（line-of-sight）；当视线被障碍物阻断时，保留阻断前的'
    '最后一个可见节点并设为新锚点。视线检测采用严格的supercover算法，确保路径不会穿越任何障碍物。'
)

add_heading_cn('1.4.2　第二阶段：角点插值平滑', level=3)
add_para(
    '对于简化后的路径，仍可能存在锐角转折。角点插值通过加权平均将转折点向平滑方向偏移。对于路径中的每个非端点'
    '节点B（前驱为A，后继为C），计算平滑点M = ((A+2B+C)/4)。平滑点M仅在不位于障碍物、且与前后节点视线均无障碍'
    '时替代原节点。此设计确保平滑过程中使用的是已平滑的前驱节点，避免连续平滑导致的路径非法问题。'
)

add_heading_cn('1.5　算法收敛性分析', level=2)
add_para(
    '本文改进算法的启发函数为α(n)·h_oct(n)。由于α(n)≥1且h_oct(n)为可接受启发函数，当α(n)>1时，改进后的启发函数'
    '可能超过真实代价，算法不再保证最优性，但能保证找到可行路径。设最优路径代价为C*，则RDA*找到的路径代价C满足'
    'C≤(1+β)·C*，即路径代价的上界为最优代价的(1+β)倍。当β=0.3时，理论最大路径代价增加率为30%，而实际实验中'
    '路径长度仅增加3.6%（通过两阶段平滑补偿），远优于理论上界。'
)

add_heading_cn('1.6　时间复杂度分析', level=2)
add_para(
    '设地图大小为H×W，路径节点数为n。RDA*算法的各阶段复杂度如下：积分图构建O(H×W)，仅执行一次；逐节点障碍率'
    '查询O(1)；A*搜索O(N\'·logN\')，其中N\'为改进后的扩展节点数，显著小于传统A*的N；路径简化O(n·L)，其中L为视线'
    '检测的平均扫描长度；角点平滑O(n)。整体时间复杂度与传统A*同阶，但因扩展节点大幅减少而实际运行更快。'
)

# ════════════════════════════════════════
# 2 实验设计
# ════════════════════════════════════════
add_heading_cn('2　实验设计', level=1)
add_heading_cn('2.1　实验环境与数据集', level=2)
add_para(
    '实验在Ubuntu 24.04操作系统下进行，编程语言为Python 3.12。本文采用MovingAI公开基准数据集[14]进行实验验证。'
    '该数据集被广泛应用于路径规划算法的性能评估，每张地图配有标准场景文件，包含大量预设的起终点对及其最优路径'
    '长度。从数据集中选取15张地图，覆盖三种类型（DAO、Street、WC3）和五个障碍率梯度（18%−65%），具体如表1所示。'
)

add_table_caption('表1　实验地图信息', 'Table 1　Information of experimental maps')
add_simple_table(
    ['类型', '地图名称', '尺寸', '障碍率/%'],
    [
        ['DAO', 'brc503d', '257×320', '18.1'],
        ['DAO', 'orz100d', '395×412', '38.8'],
        ['DAO', 'den502d', '251×211', '48.6'],
        ['DAO', 'arena2', '209×281', '58.6'],
        ['DAO', 'ost003d', '194×194', '64.9'],
        ['Street', 'Berlin_0_512', '512×512', '25.0'],
        ['Street', 'London_0_512', '512×512', '25.0'],
        ['Street', 'Moscow_0_512', '512×512', '25.0'],
        ['Street', 'Paris_0_512', '512×512', '25.0'],
        ['Street', 'Shanghai_0_512', '512×512', '25.0'],
        ['WC3', 'dustwallowkeys', '512×512', '31.5'],
        ['WC3', 'blastedlands', '512×512', '49.9'],
        ['WC3', 'bloodvenomfalls', '512×512', '55.7'],
        ['WC3', 'darkforest', '512×512', '61.9'],
        ['WC3', 'battleground', '512×512', '64.8'],
    ]
)

add_heading_cn('2.2　对比算法', level=2)
add_para(
    '实验设置7种算法进行对比，包括4种基线算法、本文算法及2种消融变体，具体如表2所示。'
)

add_table_caption('表2　对比算法设置', 'Table 2　Configuration of compared algorithms')
add_simple_table(
    ['编号', '算法名称', '启发函数', '权重', '平滑'],
    [
        ['G1', 'Dijkstra', 'h=0', '—', '无'],
        ['G2', 'A*(欧氏)', '欧氏距离', '1.0', '无'],
        ['G3', 'A*(Octile)', 'Octile距离', '1.0', '无'],
        ['G4', 'WA*(1.2)', 'Octile距离', '1.2固定', '无'],
        ['G5', 'RDA*（本文）', 'Octile距离', 'α(n)自适应', '两阶段'],
        ['G6', 'NoAdaptive', 'Octile距离', '1.0', '两阶段'],
        ['G7', 'NoSmooth', 'Octile距离', 'α(n)自适应', '无'],
    ]
)

add_heading_cn('2.3　评价指标与统计方法', level=2)
add_para(
    '采用路径长度（PL）、转弯次数（TC）、扩展节点数（EXP）、搜索时间（search_ms）、总时间（total_ms）和碰撞检测（CF）'
    '6项指标评价算法性能。对每张地图取20条任务的均值作为该地图的代表值，得到n=15个样本对。采用Wilcoxon符号秩'
    '检验（双侧）进行配对比较，使用BH方法进行FDR校正。同时报告Cohen\'s d效应量和Win/Tie/Loss统计。'
)

# ════════════════════════════════════════
# 3 实验结果与分析
# ════════════════════════════════════════
add_heading_cn('3　实验结果与分析', level=1)
add_heading_cn('3.1　参数选择', level=2)
add_para(
    '在30%−50%分位的调参集上，对β∈{0.1, 0.2, 0.3, 0.4, 0.5}进行网格搜索，结果如表3所示。β=0.3在PL增加率<5%的'
    '约束下实现了最大的EXP减少率（52.2%），且PL增加可通过两阶段平滑进一步补偿，因此选定β*=0.3。'
)

add_table_caption('表3　β参数调优结果', 'Table 3　Results of β parameter tuning')
add_simple_table(
    ['β', '平均EXP减少率/%', '平均PL增加率/%', '是否可行'],
    [
        ['0.1', '35.5', '1.4', '是'],
        ['0.2', '46.4', '2.9', '是'],
        ['0.3', '52.2', '4.5', '是（选定）'],
        ['0.4', '56.0', '5.7', '否（PL>5%）'],
        ['0.5', '59.5', '7.1', '否（PL>5%）'],
    ]
)

# β敏感性图
add_image('figures/chart_beta_sensitivity.png', 12)
add_fig_caption('图1　β参数敏感性分析', 'Fig.1　β parameter sensitivity analysis')

add_heading_cn('3.2　主实验结果', level=2)
add_para(
    '主实验在15张地图上以20条任务×7种算法执行，共计2100次搜索。所有搜索均成功完成，超时率为0，碰撞合法率为100%。'
    '全局均值结果如表4所示。'
)

add_table_caption('表4　主实验全局均值结果', 'Table 4　Global mean results of main experiment')
add_simple_table(
    ['算法', 'PL', 'TC', 'EXP', 'search_ms', 'total_ms'],
    [
        ['Dijkstra', '312.94', '19.6', '81568', '559.21', '559.21'],
        ['A*(欧氏)', '312.94', '51.1', '17019', '126.31', '126.31'],
        ['A*(Octile)', '312.94', '22.5', '13331', '113.43', '113.43'],
        ['WA*(1.2)', '318.53', '37.5', '7171', '60.49', '60.49'],
        ['RDA*（本文）', '301.76', '12.1', '6550', '80.67', '86.04'],
        ['NoAdaptive', '299.53', '8.7', '13331', '113.22', '119.08'],
        ['NoSmooth', '328.41', '57.9', '6550', '80.52', '80.52'],
    ]
)

add_para(
    '由表4可知，与A*(欧氏)相比，本文RDA*算法在所有核心指标上均取得显著改善：路径长度缩短3.6%，转弯次数减少76.4%，'
    '扩展节点减少61.5%，总运行时间缩短31.9%。与WA*(1.2)相比，RDA*在路径质量上具有明显优势（PL 301.76 vs 318.53，'
    'TC 12.1 vs 37.5），同时扩展节点更少（6550 vs 7171）。'
)

# 柱状图
add_image('figures/chart_expanded_nodes.png', 14)
add_fig_caption('图2　15张地图扩展节点数对比', 'Fig.2　Comparison of expanded nodes across 15 maps')

add_image('figures/chart_path_length.png', 14)
add_fig_caption('图3　15张地图路径长度对比', 'Fig.3　Comparison of path length across 15 maps')

add_image('figures/chart_turn_count.png', 14)
add_fig_caption('图4　15张地图转弯次数对比', 'Fig.4　Comparison of turn count across 15 maps')

# 7算法面板图
add_image('figures/panel_7algo_publication.png', 14)
add_fig_caption('图5　7种算法路线对比（bloodvenomfalls）', 'Fig.5　Route comparison of 7 algorithms (bloodvenomfalls)')

add_heading_cn('3.3　统计检验', level=2)
add_para('对RDA*与各基线算法的配对比较进行Wilcoxon符号秩检验，结果如表5所示。')

add_table_caption('表5　RDA*与A*(欧氏)的统计检验结果', 'Table 5　Statistical test results of RDA* vs A*(Euclidean)')
add_simple_table(
    ['指标', '变化率/%', 'p_raw', 'q_bh', '显著性', 'W/T/L'],
    [
        ['PL', '−3.6', '0.000061', '0.000105', '***', '15/0/0'],
        ['TC', '−76.4', '0.000061', '0.000105', '***', '15/0/0'],
        ['EXP', '−61.5', '0.000061', '0.000105', '***', '15/0/0'],
        ['total_ms', '−31.9', '0.015', '0.020', '*', '12/0/3'],
        ['search_ms', '−36.1', '0.012', '0.018', '*', '12/0/3'],
    ]
)

add_para(
    'PL、TC、EXP三项核心指标的q_bh值均远小于0.001（极显著），W/T/L=15/0/0，即RDA*在所有15张地图上均优于A*(欧氏)。'
)

# 路线叠加对比图
add_image('figures/compare_bloodvenomfalls.png', 14)
add_fig_caption('图6　路线叠加对比（bloodvenomfalls，obs=55.7%）',
                'Fig.6　Route overlay comparison (bloodvenomfalls, obs=55.7%)')

add_image('figures/compare_Berlin_0_512.png', 14)
add_fig_caption('图7　路线叠加对比（Berlin，obs=25.0%）',
                'Fig.7　Route overlay comparison (Berlin, obs=25.0%)')

add_heading_cn('3.4　消融分析', level=2)
add_para(
    '为验证自适应权重和路径平滑两个模块各自的贡献，对RDA*与两种消融变体进行对比分析，结果如表6所示。'
)

add_table_caption('表6　消融实验结果', 'Table 6　Results of ablation study')
add_simple_table(
    ['对比', '指标', '变化率/%', 'p_raw', '显著性'],
    [
        ['RDA* vs NoAdaptive', 'EXP', '−50.9', '0.000061', '***'],
        ['RDA* vs NoAdaptive', 'total_ms', '−27.7', '0.030', '*'],
        ['RDA* vs NoAdaptive', 'PL', '+0.7', '0.000061', '***'],
        ['RDA* vs NoSmooth', 'PL', '−8.1', '0.000061', '***'],
        ['RDA* vs NoSmooth', 'TC', '−79.2', '0.000061', '***'],
        ['RDA* vs NoSmooth', 'EXP', '0', '—', '=='],
    ]
)

add_para(
    '消融结果表明：（1）自适应权重模块的主要贡献在于搜索效率——使扩展节点减少50.9%，搜索时间缩短27.7%，代价是路径'
    '长度微增0.7%；（2）路径平滑模块的主要贡献在于路径质量——使路径长度缩短8.1%，转弯次数减少79.2%，而不改变搜索'
    '阶段的行为。两个模块功能正交、互为补充。'
)

add_image('figures/chart_ablation.png', 12)
add_fig_caption('图8　消融实验对比', 'Fig.8　Ablation study comparison')

add_heading_cn('3.5　困难场景验证', level=2)
add_para(
    '为验证算法在困难场景下的鲁棒性，在相同的15张地图上使用最长路径任务（平均最优路径≈626步）进行补充实验。结果'
    '如表7所示。补充实验的改进趋势与主实验完全一致，三项核心指标的p值均小于0.001，W/T/L均为15/0/0，说明RDA*算法'
    '在困难场景下仍具有良好的鲁棒性。'
)

add_table_caption('表7　主实验与补充实验对比', 'Table 7　Comparison of main and supplementary experiments')
add_simple_table(
    ['指标', '主实验(PL≈316)', '补充实验(PL≈626)'],
    [
        ['PL变化率', '−3.6% ***', '−2.8% ***'],
        ['TC变化率', '−76.4% ***', '−71.9% ***'],
        ['EXP变化率', '−61.5% ***', '−58.4% ***'],
    ]
)

# ════════════════════════════════════════
# 4 结论
# ════════════════════════════════════════
add_heading_cn('4　结论', level=1)

add_para(
    '本文针对传统A*算法在复杂栅格环境中搜索效率受限、路径冗余转折多、启发函数权重无法自适应调整等问题，提出了'
    '残差驱动自适应加权A*路径规划算法（RDA*），通过融合局部障碍物密度感知与两阶段路径平滑策略，实现了搜索效率与'
    '路径质量的协同优化。'
)

add_para(
    '在搜索策略方面，本文以Octile距离替代传统欧氏距离作为启发函数，更准确地估计八连通栅格中的实际路径代价。在此'
    '基础上，引入局部障碍物密度概念，利用积分图实现O(1)复杂度的逐节点障碍率查询，设计自适应权重函数α(n)=1+β(1−'
    'ρ_local(n))，使算法能够根据节点周围的环境特征动态调整搜索策略——在空旷区域增大启发权重以加速搜索，在障碍物'
    '密集区域降低权重以保守寻优。通过在调参集上的网格搜索确定β*=0.3，在路径长度增加率控制在5%以内的前提下，实现'
    '了52.2%的扩展节点减少率。'
)

add_para(
    '在路径优化方面，本文提出两阶段平滑策略：第一阶段基于严格supercover视线检测进行路径简化，去除冗余中间节点；'
    '第二阶段通过角点加权插值平滑路径转折，降低路径曲率变化。两个阶段均内置碰撞合法性检查，确保平滑后的路径不'
    '穿越任何障碍物。该策略有效补偿了加权搜索带来的路径次优性，使最终路径长度反而优于标准A*算法。'
)

add_para(
    '实验验证方面，本文在MovingAI公开基准数据集的15张地图上进行了大规模系统性实验。主实验以7种算法×20条任务共计'
    '2100次搜索为基础，辅以Wilcoxon符号秩检验和BH-FDR校正进行统计分析。结果表明，与A*(欧氏)相比，RDA*的路径长度'
    '缩短3.6%，转弯次数减少76.4%，扩展节点减少61.5%，总运行时间缩短31.9%，四项指标均达到统计显著水平。消融实验'
    '证实了自适应权重模块与路径平滑模块各自具有独立且显著的贡献：前者使扩展节点减半，后者使路径长度缩短8.1%、'
    '转弯次数减少79.2%，两个模块功能正交、互为补充。补充实验在长路径场景（平均最优路径≈626步）下进一步验证了算法'
    '的鲁棒性，改进趋势与主实验完全一致。'
)

add_para(
    '同时，本文也存在以下局限：（1）在障碍物分布均匀的环境中，局部障碍率的空间差异较小，自适应权重退化为近似固定值，'
    '算法优势减弱；（2）当前实验仅在仿真环境中进行，尚未在实际机器人平台上验证算法的实时性和可行性。'
)

add_para(
    '未来的研究可从以下方向展开：引入β参数的自动衰减机制，使算法在障碍物分布均匀的环境中也能保持竞争力；增加JPS、'
    'Theta*等高级路径规划算法作为对比基线，进一步验证算法的相对优势；开展Gazebo仿真与真实机器人平台的实验，验证'
    '算法在动态环境中的实际应用效果。'
)

# ════════════════════════════════════════
# 参考文献
# ════════════════════════════════════════
add_heading_cn('参考文献：', level=1)

refs = [
    '[1] 朱大奇, 颜明重. 移动机器人路径规划技术综述[J]. 控制与决策, 2010, 25(7): 961-967.',
    '[2] 王旭, 朱其新, 朱永红. 面向二维移动机器人的路径规划算法综述[J]. 计算机工程与应用, 2023, 59(20): 51-66.',
    '[3] 毛建旭, 贺振宇, 王耀南, 等. 电力巡检机器人路径规划技术及应用综述[J]. 控制与决策, 2023, 38(11): 3009-3024.',
    '[4] STURTEVANT N R. Benchmarks for grid-based pathfinding[J]. IEEE Transactions on Computational Intelligence and AI in Games, 2012, 4(2): 144-148.',
    '[5] HART P E, NILSSON N J, RAPHAEL B. A formal basis for the heuristic determination of minimum cost paths[J]. IEEE Transactions on Systems Science and Cybernetics, 1968, 4(2): 100-107.',
    '[6] ALGFOOR Z A, SUNAR M S, ABDULLAH A. A new weighted pathfinding algorithms to reduce the search time on grid maps[J]. Expert Systems with Applications, 2017, 71: 319-331.',
    '[7] 冯志乾, 王欣, 吴迪. 基于改进A*和DWA融合的移动机器人路径规划[J]. 计算机应用与软件, 2026, 43(2): 340-346.',
    '[8] 高建京, 崔明月. 融合改进A*算法和动态窗口法的移动机器人路径规划研究[J/OL]. 控制工程, 2026.',
    '[9] 毕竟, 刘俊. 结合DC-A*与FE-DWA的巡检机器人路径规划方法[J]. 计算机工程与应用, 2026, 62(3): 334-346.',
    '[10] 赵江, 张岩, 马泽文. AGV路径规划A星算法的改进与验证[J]. 计算机工程与应用, 2018, 54(21): 217-223.',
    '[11] 冯泽鹏, 李宗刚, 夏广庆, 等. 改进A*与APF的移动机器人路径规划算法研究[J]. 计算机工程与应用, 2025, 61(20): 132-145.',
    '[12] 赵晓, 王铮, 黄程侃. 基于改进A*算法的移动机器人路径规划[J]. 机器人, 2018, 40(6): 903-910.',
    '[13] 吴鹏, 桑成军, 陆忠华, 等. 基于改进A*算法的移动机器人路径规划研究[J]. 计算机工程与应用, 2019, 55(21): 227-233.',
    '[14] STURTEVANT N R. Benchmarks for grid-based pathfinding[J]. IEEE Transactions on Computational Intelligence and AI in Games, 2012, 4(2): 144-148.',
]

for ref in refs:
    add_para(ref, size=9, indent=False)

# ════════════════════════════════════════
# 保存
# ════════════════════════════════════════
out_path = '/mnt/user-data/outputs/RDA_star_论文初稿.docx'
doc.save(out_path)
print(f"✅ 论文保存至: {out_path}")
